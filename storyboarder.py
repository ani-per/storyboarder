from pathlib import Path  # Filepaths
from typing import Tuple  # Argument / output type checking
from re import sub, search, findall, split # String operations
from itertools import compress

import pandas as pd  # DataFrames
import docx as docx  # Word documents - https://github.com/python-openxml/python-docx
from haggis.files.docx import (
    list_number,
)  # Misc Word document utils - https://gitlab.com/madphysicist/haggis
from tmdbv3api import (
    TMDb,
    Movie,
    Search,
)  # TMDB API interface - https://github.com/AnthonyBloomer/tmdbv3api

from subprocess import call
from platform import system

tmdb = TMDb()
tmdb.api_key = "57989a14e4e5073fc2d332b37967de77"  # Ani Perumalla's personal key


# https://stackoverflow.com/a/40319071
def _copy(self, target):
    from shutil import copy as sh_copy

    assert self.is_file()
    sh_copy(str(self), str(target))


Path.copy = _copy


# https://stackoverflow.com/a/38234962
def make_curly(str: str) -> str:
    return sub(
        r"(\s|^)\'(.*?)\'(\s|$)", r"\1‘\2’\3", sub(r"\"(.*?)\"", r"“\1”", str)
    ).replace("'", "’")


def style_doc(tmpl):
    tmpl.styles["Title"].font.name = "Times New Roman"
    tmpl.styles["Heading 1"].font.name = "Times New Roman"
    tmpl.styles["Normal"].font.name = "Times New Roman"
    pass


def get_main_ans(ans_raw: str) -> Tuple[str, str, list]:
    ans_split = ans_raw.split(" [")
    if len(ans_split) > 1:
        alt_ans = search("\[(.*?)\]", "[" + ans_split[1]).group(1).split("; ")
    else:
        alt_ans = []
    main_ans_split = ans_split[0].split(" (")
    if len(main_ans_split) > 1:
        pg_ans = f" ({main_ans_split[1]}"
    else:
        pg_ans = ""
    main_ans = main_ans_split[0]

    return (main_ans, pg_ans, alt_ans)


def write_answerline(
    ans_par: docx.text.paragraph.Paragraph,
    main_ans: str,
    pg_ans: str,
    alt_ans: list,
    ans_type: str,
):
    """Write the database answerline to a given paragraph in a document.

    Args:
        ans_par (docx.text.paragraph.Paragraph): The paragraph to which the answerline should be written.
        main_ans (str): The main answerline.
        pg_ans (str): The pronunciation guide for the answerline.
        alt_ans (str): The list of alternate answerlines.
        ans_type (str): The class of the answerline in the database (e.g. "Film", "Creator", "Director", "Crew", "Figure", "Surname", "Misc").
    """
    articles = (
        "A ",
        "The ",
        "a ",
        "the ",
        "Le ",
        "le ",
        "La ",
        "la ",
        "El ",
        "el ",
        "Il ",
        "il ",
        "Dir ",
        "dir ",
    )

    # Style the main answerline first
    if ans_type in [
        "Creator",
        "Director",
        "Crew",
        "Figure",
        "Surname",
    ]:  # If it's a person, format their surname
        main_names = main_ans.split(" ")
        n_main_names = len(main_names)
        main_runs = n_main_names * [None]
        for i in range(n_main_names):
            main_runs[i] = ans_par.add_run(main_names[i])
            if i == (n_main_names - 1):
                main_runs[i].bold = True
                main_runs[i].underline = True
            else:
                ans_par.add_run(" ")
    else:
        if main_ans.startswith(articles):  # Don't format the article
            main_words = main_ans.split(" ")
            article_run = ans_par.add_run(main_words[0] + " ")  # Unformatted article
            main_run = ans_par.add_run(" ".join(main_words[1:]))  # Rest of answerline
        else:  # Just print the answerline as-is
            article_run = ans_par.add_run("")
            main_run = ans_par.add_run(main_ans)
        if ans_type == "Film":  # If it's a film, italicize it
            article_run.italic = True
            main_run.italic = True

        main_run.bold = True
        main_run.underline = True

    # Add the pronunciation guide
    ans_par.add_run(pg_ans)

    # Style the alt answerlines if they exist
    n_alt_ans = len(alt_ans)
    if n_alt_ans > 0:
        for i in range(n_alt_ans):
            for directive in ["or ", "accept ", "prompt on ", "reject "]:
                if alt_ans[i].startswith(directive):
                    if i == 0:
                        ans_par.add_run(" [")
                    ans_par.add_run(directive)
                    ans_val = split("^" + directive, alt_ans[i])[-1]
                    if ans_type in ["Creator", "Director", "Crew", "Figure", "Surname"]:
                        alt_names = ans_val.split(" ")
                        n_alt_names = len(alt_names)
                        alt_name_runs = n_alt_names * [None]
                        for j in range(n_alt_names):
                            alt_name_runs[j] = ans_par.add_run(alt_names[j])
                            if j == (n_alt_names - 1):
                                if not directive.startswith("reject"):
                                    alt_name_runs[j].underline = True
                                    if not directive.startswith("prompt"):
                                        alt_name_runs[j].bold = True
                            else:
                                ans_par.add_run(" ")
                    else:
                        if ans_val.startswith(articles):  # Don't format the article
                            alt_words = ans_val.split(" ")
                            article_run = ans_par.add_run(
                                alt_words[0] + " "
                            )  # Unformatted article
                            alt_run = ans_par.add_run(
                                " ".join(alt_words[1:])
                            )  # Rest of answerline
                        else:  # Just print the answerline as-is
                            article_run = ans_par.add_run("")
                            alt_run = ans_par.add_run(ans_val)
                        if ans_type == "Film":  # If it's a film, italicize it
                            article_run.italic = True
                            alt_run.italic = True
                        if not directive.startswith("reject"):
                            alt_run.underline = True
                            if not directive.startswith("prompt"):
                                alt_run.bold = True
                    if i == (n_alt_ans - 1):
                        ans_par.add_run("]")
                    else:
                        ans_par.add_run("; ")


def storyboard(
    set_name: str,
    db_dir: Path,
    db_path: Path,
    hybrid: bool = False,
    raw_string: str = "_raw",
    src_dir: Path = None,
    dest_dir: Path = None,
    split_docs: bool = False,
    tags: bool = True,
    force_end: bool = True,
    verbose: bool = False,
    try_open: bool = False,
    n_visual_questions: int = 10,
    n_total_questions: int = 20,
):
    """Create the visual answerline document, and hybrid packets along the way if desired.

    Args:
        set_name (str): Name of the set/packet
        db_dir (Path): Directory where the answerline database CSV is stored (and where the visual answerline document will be generated)
        db_path (Path): Filepath of the answerline database CSV
        hybrid (bool, optional): Is the set a hybrid visual-written tournament? Defaults to True.
        raw_string (str, optional): Suffix to identify the written packets. All written packets must end with this prefix for Storyboarder to identify them. Defaults to "_raw".
        src_dir (Path, optional): Directory where the written packets are stored (and where the hybrid packets will be generated). Defaults to `db_dir`.
        dest_dir (Path, optional): Directory where the hybrid packets will be generated. Defaults to `src_dir`.
        split_docs (bool, optional): Should the answerline documents be split? Defaults to False.
        tags (bool, optional): Should the hybrid packets have author tags for the visual questions? Defaults to False.
        force_end (bool, optional): Should the written equivalent of the visual slides end with the penultimate number for MODAQ use, or should they end with the final number? Defaults to True.
        verbose (bool, optional): Print progress. Defaults to False.
        try_open (bool, optional): Try to open the generated answerlines document. Defaults to False.
        n_visual_questions (int, optional): The number of visual questions in each packet, if hybrid. Defaults to 10.
        n_total_questions (int, optional): The number of total questions in each packet, including written and visual if hybrid. Defaults to 20.
    """
    media = ["Film", "Music Video", "Video", "Television"]
    people = ["Creator", "Director"]

    movie = Movie()
    search = Search()

    ans_db = pd.read_csv(db_path).convert_dtypes()

    if src_dir is None:
        src_dir = db_dir

    set_slug = set_name.title().replace(" ", "-")

    packet_names = list(ans_db["Packet"].unique())
    n_packet = len(packet_names)

    templates = n_packet * [None]
    documents = n_packet * [None]
    packets = n_packet * [None]
    answers = n_packet * [ans_db["Number"].max() * [None]]
    slides = n_packet * [ans_db["Number"].max() * [[None]]]

    for i in range(n_packet):
        if split_docs or (not split_docs and (i == 0)):
            templates[i] = docx.Document()
            if split_docs:
                documents[i] = (
                    db_dir / f"{set_slug}_Answers-raw_{packet_names[i].zfill(2)}.docx"
                )
            elif not split_docs and (i == 0):
                documents[i] = db_dir / f"{set_slug}_Answers-raw.docx"
        elif (not split_docs) and (i > 0):
            templates[i] = templates[0]
            documents[i] = documents[0]

    # Use curly quotes
    for col in ["Answerline", "Source", "Creator", "Director", "Notes"]:
        if col in ans_db.columns:
            ans_db[col] = ans_db[col].apply(
                lambda s: make_curly(s) if pd.notna(s) else s
            )

    # Prepare the hybrid packet generation, if configured
    make_hybrid = False
    if hybrid and src_dir.exists():
        written_packets = sorted(src_dir.glob(f"**/[!~]?*{raw_string}.docx"))
        if len(list(written_packets)) == 0:
            print("There aren't any written packets in the source folder.")
            return
        if dest_dir is None:
            dest_dir = src_dir
        if len(list(written_packets)) == n_packet:
            make_hybrid = True
            hybrid_packets = n_packet * [None]
            hybrid_answers = n_packet * [ans_db["Number"].max() * [None]]
        else:
            print(
                "Number of placeholder packets doesn't match the number of written packets."
            )

    for i in range(n_packet):  # Loop over packets
        style_doc(templates[i])
        if (split_docs) or ((not split_docs) and (i == 0)):
            templates[i].add_heading(f"{set_name} - Visual Answerlines", level=0)
        elif (not split_docs) and (i > 0):
            templates[i].add_page_break()
        packet_header = f"Packet {packet_names[i]}"
        packets[i] = templates[i].add_heading(packet_header, level=1)

        if verbose:
            print(packet_header)

        # Filter the database to the current packet
        packet_db = ans_db[ans_db["Packet"] == packet_names[i]]

        # If hybrid, make the current hybrid packet by appending to the corresponding written packet
        if make_hybrid:
            hybrid_packets[i] = dest_dir / f"{set_slug}_{packet_names[i].zfill(2)}.docx"
            if hybrid_packets[i].exists():
                hybrid_packets[i].unlink()
            Path.copy(written_packets[i], hybrid_packets[i])
            hybrid_docx = docx.Document(hybrid_packets[i])
            # Calculate the number of already-written questions in the packet
            # This compiles all the numbers in the document that begin a new line and are succeeded by a period, then takes the maximum of those numbers
            # https://stackoverflow.com/questions/952914/how-do-i-make-a-flat-list-out-of-a-list-of-lists#comment123215183_952952
            detect_written = [
                int(s)
                for s in [
                    leaf
                    for tree in [
                        findall("(^\d+)\.+", s)
                        for s in [p.text for p in hybrid_docx.paragraphs]
                    ]
                    for leaf in tree
                    if leaf
                ]
            ]
            if len(detect_written) > 0:
                n_written = max(detect_written)
                if n_written > (n_total_questions - n_visual_questions):
                    print(
                        f"Skipping packet {i + 1} - there are already {n_written} questions, which means there isn't space for {n_visual_questions} more questions in packets of {n_total_questions} questions."
                    )
                    continue
            else:
                n_written = 0

        for j in range(packet_db["Number"].max()):  # Loop over questions
            # Filter the database to the current question
            q_db = packet_db[packet_db["Number"] == (j + 1)]
            n_slide = q_db.shape[0]

            # Only process if the answerline is not empty
            if pd.notna(q_db.iloc[0]["Answerline"]):
                if verbose:
                    print(f"{j + 1}: " + q_db.iloc[0]["Answerline"])

                # Write the answerline
                answers[i][j] = templates[i].add_paragraph("", style="List Number")
                main_ans, pg_ans, alt_ans = get_main_ans(q_db.iloc[0]["Answerline"])
                write_answerline(
                    answers[i][j],
                    main_ans,
                    pg_ans,
                    alt_ans,
                    q_db.iloc[0]["Answerline_Type"],
                )

                # If hybrid, make the placeholder question
                if make_hybrid:
                    if j > 0:
                        hybrid_docx.add_paragraph("")
                    slide_q = hybrid_docx.add_paragraph(f"{j + n_written + 1}. ")
                    slide_runs = n_slide * [None]
                    for k in range(n_slide):  # Loop over slides
                        if force_end or k < (n_slide - 1):
                            slide_runs[k] = slide_q.add_run(
                                f"{k + 1}"
                                + " "
                                * (
                                    (not force_end and k < (n_slide - 2))
                                    or (force_end and k < (n_slide - 1))
                                )
                            )
                        if k < (n_slide - 1):
                            if q_db.iloc[k]["Value"] == 20:
                                slide_runs[k].bold = True
                                slide_runs[k].underline = True
                                if q_db.iloc[k + 1]["Value"] < 20:
                                    superpower_mark = slide_q.add_run("(+)")
                                    superpower_mark.bold = True
                                    superpower_mark.underline = True
                                    slide_q.add_run(" ")
                            elif q_db.iloc[k]["Value"] == 15:
                                slide_runs[k].bold = True
                                if q_db.iloc[k + 1]["Value"] < 15:
                                    power_mark = slide_q.add_run("(*)")
                                    power_mark.bold = True
                                    slide_q.add_run(" ")

                    hybrid_answers[i][j] = hybrid_docx.add_paragraph("ANSWER: ")
                    write_answerline(
                        hybrid_answers[i][j],
                        main_ans,
                        pg_ans,
                        alt_ans,
                        q_db.iloc[0]["Answerline_Type"],
                    )

                # Print the source metadata (creators, directors)
                if (
                    q_db.iloc[0]["Answerline_Type"] == "Film"
                ):  # If it's a film, use the answerline to search
                    director = ""
                    if pd.notna(q_db.iloc[0]["Creator"]):
                        director = q_db.iloc[0]["Creator"]
                    else:
                        results = search.movies(
                            main_ans,
                            year=q_db.iloc[0]["Source_Year"]
                            if pd.notna(q_db.iloc[0]["Source_Year"])
                            else "",
                        )
                        if results["total_results"] > 0:
                            director = make_curly(", ".join(
                                [
                                    crew["name"]
                                    for crew in movie.credits(
                                        results["results"][0].id
                                    ).crew
                                    if crew["job"] == "Director"
                                ]
                            ))
                    if len(director) > 0:
                        dir_raw = f" (dir. {director})"
                    else:
                        dir_raw = ""
                    answers[i][j].add_run(dir_raw)
                    if make_hybrid:
                        hybrid_answers[i][j].add_run(dir_raw)
                else:  # Prepare the slide annotations, if it's not a film
                    if q_db.iloc[0]["Answerline_Type"] not in people:
                        # Extract the metadata for sources that are films
                        film_data = q_db[["Source", "Source_Type", "Source_Year", "Creator"]].drop_duplicates().reset_index(drop=True)

                        for k in range(len(film_data)): # Loop over films
                            if (pd.isna(film_data["Creator"][k])):
                                results = search.movies(
                                    film_data['Source'][k],
                                    year=film_data["Source_Year"][k]
                                    if pd.notna(film_data["Source_Year"][k])
                                    else "",
                                )
                                if results["total_results"] > 0:
                                    film_data.at[k, "Creator"] = make_curly(", ".join(
                                        [
                                            crew["name"]
                                            for crew in movie.credits(
                                                results["results"][0].id
                                            ).crew
                                            if crew["job"] == "Director"
                                        ]
                                    ))

                    slides[i][j] = n_slide * [None]
                    for k in range(n_slide):  # Loop over slides
                        slides[i][j][k] = templates[i].add_paragraph(
                            "", style="List Number 2"
                        )

                        # Add an annotation if there's a source listed for the current slide
                        src_exists = pd.notna(q_db.iloc[k]["Source"])
                        src_raw = q_db.iloc[k]["Source"] if src_exists else ""
                        src_run = slides[i][j][k].add_run(src_raw)
                        if (src_exists) and not (
                            q_db.iloc[k]["Source"].startswith(("'", '"', "‘", "“"))
                        ):  # Don't italicize if title's in quotes (e.g. music video)
                            src_run.italic = True

                        # If the question's not on a creator, add the director credit for the source of the current slide
                        if (
                            q_db.iloc[0]["Answerline_Type"]
                            not in people
                        ):
                            creator = ""
                            if pd.notna(q_db.iloc[k]["Creator"]):
                                creator = q_db.iloc[k]["Creator"]
                            else:
                                creator = film_data["Creator"][film_data["Source"].eq(q_db.iloc[k]["Source"]).idxmax()]

                            if pd.isna(q_db.iloc[k]["Source_Type"]) or (q_db.iloc[k]["Source_Type"] in media):
                                credit = "dir."
                            else:
                                credit = "by"
                            if (pd.notna(creator)) and (len(creator) > 0):
                                slides[i][j][k].add_run(f" ({credit} {creator})")

                        # Format the annotation as a list element
                        if k == 0:
                            list_number(
                                templates[i], slides[i][j][k], prev=None, level=0
                            )
                        else:
                            list_number(
                                templates[i], slides[i][j][k], prev=slides[i][j][k - 1]
                            )

                    # If hybrid, write the sources in the visual question as a note
                    if make_hybrid:
                        hybrid_answers[i][j].add_run(" (Sources: ")
                        if q_db.iloc[0]["Answerline_Type"] in people:
                            films = q_db["Source"][q_db["Source"].notnull()].unique()
                            for k in range(len(films)):  # Loop over films
                                if k > 0:
                                    hybrid_answers[i][j].add_run("; ")
                                hybrid_answers[i][j].add_run(films[k]).italic = True
                        else:
                            directors = film_data["Creator"].unique()
                            for k in range(len(directors)):  # Loop over directors
                                srcs = film_data[film_data["Creator"] == directors[k]]["Source"].unique()
                                if k > 0:
                                    hybrid_answers[i][j].add_run("; ")
                                for l in range(len(srcs)):
                                    if l > 0:
                                        hybrid_answers[i][j].add_run(", ")
                                    if srcs[
                                        l
                                    ].startswith(
                                        ("'", '"', "‘", "“")
                                    ):  # Don't italicize if title's in quotes (e.g. music video)
                                        hybrid_answers[i][j].add_run(srcs[l])
                                    else:
                                        hybrid_answers[i][j].add_run(
                                            srcs[l]
                                        ).italic = True
                                if pd.isna(film_data.iloc[k]["Source_Type"]) or (film_data.iloc[k]["Source_Type"] in media):
                                    credit = "- dir."
                                else:
                                    credit = "by"
                                hybrid_answers[i][j].add_run(f" {credit} {directors[k]}")
                        hybrid_answers[i][j].add_run(")")

                # If hybrid, write the author tag
                if make_hybrid and tags:
                    if pd.notna(q_db.iloc[0]["Author"]):
                        hybrid_docx.add_paragraph(f"<{q_db.iloc[0]['Author']}, Visual>")
                    else:
                        hybrid_docx.add_paragraph(f"<Visual>")

                # If there are notes, write them
                if pd.notna(q_db.iloc[0]["Notes"]):
                    notes_raw = f" ({q_db.iloc[0]['Notes']})"
                    answers[i][j].add_run(notes_raw)

                if j == 0:
                    list_number(templates[i], answers[i][j], prev=None, level=0)
                else:
                    list_number(templates[i], answers[i][j], prev=answers[i][j - 1])

        if make_hybrid:
            hybrid_docx.save(hybrid_packets[i])
        # Write the document
        templates[i].save(documents[i])

    if not split_docs and try_open:
        # https://stackoverflow.com/a/435669
        if system() == "Darwin":  # MacOS
            call(("open", documents[0]))
        elif system() == "Windows":  # Windows
            from os import startfile

            startfile(documents[0])
        else:  # Linux variants
            call(("xdg-open", documents[0]))
